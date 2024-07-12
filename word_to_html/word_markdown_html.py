import fitz  # PyMuPDF
import docx  # python-docx
import re
from jinja2 import Template
import os
from docx import Document
import subprocess

def read_pdf_file(file_path):
    try:
        doc = fitz.open(file_path)
        content = []
        list_stack = []  # Stack to handle nested lists
        is_contents_section = False
        contents_table = []
        paragraph_buffer = []  # Buffer to collect related paragraphs

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")

            print(f"Processing page {page_num + 1}")  # Debugging statement

            for line in text.split('\n'):
                line = line.strip()
                # Remove unwanted links and anchors
                line = re.sub(r'\[.*?\]\{#.*?\}', '', line)
                if line:
                    if re.search(r'\bCONTENTS?\b', line.upper()):
                        is_contents_section = True
                        contents_table.append("<table>")
                        continue

                    if is_contents_section:
                        if not re.match(r'^\d+\.\s', line):
                            is_contents_section = False
                            contents_table.append("</table>")
                            content.extend(contents_table)
                            contents_table = []

                    if is_contents_section:
                        formatted_line = format_contents_line(line)
                        contents_table.append(formatted_line)
                    else:
                        if re.match(r'^[0-9]+\.\s', line):
                            # Line starts with a number followed by a period and a space
                            line = f"<b>{line}</b>"
                        paragraph_buffer.append(line)

                elif paragraph_buffer:
                    # End of a paragraph block
                    formatted_paragraph = format_paragraph(paragraph_buffer, list_stack)
                    if formatted_paragraph.strip():  # Ensure non-empty paragraph
                        print(f"Formatted paragraph: {formatted_paragraph}")  # Debugging statement
                        content.append(formatted_paragraph)
                    paragraph_buffer = []

        # Handle any remaining paragraphs
        if paragraph_buffer:
            formatted_paragraph = format_paragraph(paragraph_buffer, list_stack)
            if formatted_paragraph.strip():  # Ensure non-empty paragraph
                content.append(formatted_paragraph)

        # Close any remaining open lists
        while list_stack:
            list_type = list_stack.pop()
            content.append(f'</{list_type}>')

        return content
    except Exception as e:
        print(f"An error occurred: {e}")
        return []

def read_docx_file(file_path):
    try:
        doc = Document(file_path)
        content = []
        list_stack = []  # Stack to handle nested lists
        is_contents_section = False
        contents_table = []
        paragraph_buffer = []  # Buffer to collect related paragraphs

        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()

            # Remove unwanted links and anchors
            line = re.sub(r'\[.*?\]\{#.*?\}', '', line)

            # Remove content inside curly braces {}
            line = re.sub(r'\{.*?\}', '', line)

            if line:
                if re.search(r'\bCONTENTS?\b', line.upper()):
                    is_contents_section = True
                    contents_table.append("<table>")
                    continue

                if is_contents_section:
                    if not re.match(r'^\d+\.\s', line):
                        is_contents_section = False
                        contents_table.append("</table>")
                        content.extend(contents_table)
                        contents_table = []

                if is_contents_section:
                    formatted_line = format_contents_line(line)
                    contents_table.append(formatted_line)
                else:
                    if re.match(r'^[0-9]+\.\s', line):
                        # Line starts with a number followed by a period and a space
                        line = f"<b>{line}</b>"
                    paragraph_buffer.append(line)

            elif paragraph_buffer:
                # End of a paragraph block
                formatted_paragraph = format_paragraph(paragraph_buffer, list_stack)
                if formatted_paragraph.strip():  # Ensure non-empty paragraph
                    print(f"Formatted paragraph: {formatted_paragraph}")  # Debugging statement
                    content.append(formatted_paragraph)
                paragraph_buffer = []

        # Handle any remaining paragraphs
        if paragraph_buffer:
            formatted_paragraph = format_paragraph(paragraph_buffer, list_stack)
            if formatted_paragraph.strip():  # Ensure non-empty paragraph
                content.append(formatted_paragraph)

        # Close any remaining open lists
        while list_stack:
            list_type = list_stack.pop()
            content.append(f'</{list_type}>')

        return content
    except Exception as e:
        print(f"An error occurred: {e}")
        return []

def format_contents_line(line):
    parts = re.split(r'(.{2,})', line)
    if len(parts) >= 3:
        title = parts[0].strip()
        page = parts[-1].strip()
        return f'<tr><td>{title}</td><td style="text-align: right;">{page}</td></tr>'
    return f'<tr><td colspan="2">{line}</td></tr>'

def format_paragraph(paragraphs, list_stack):
    formatted_text = "<p>" + " ".join(paragraphs) + "</p>"
    list_type = None

    # Handling bullets and numbering
    if re.match(r'^\(\d+\)', paragraphs[0]):
        list_type = 'ol'

    if list_type:
        if not list_stack or list_stack[-1] != list_type:
            if list_stack:
                closing_list_type = list_stack.pop()
                formatted_text = f'</{closing_list_type}>{formatted_text}'
            list_stack.append(list_type)
            formatted_text = f'<{list_type}><li>{formatted_text}</li>'
        else:
            formatted_text = f'<li>{formatted_text}</li>'
    else:
        if list_stack:
            closing_list_type = list_stack.pop()
            formatted_text = f'</{closing_list_type}>{formatted_text}'

    # Ensure the numbering format (e.g., "(1)", "(2)") is preserved
    numbered_list_match = re.match(r'^\(\d+\)', paragraphs[0])
    if numbered_list_match:
        formatted_text = f'{numbered_list_match.group(0)} {formatted_text[len(numbered_list_match.group(0)):].strip()}'

    return formatted_text.strip()

def create_jinja2_template(content, template_path, filename):
    try:
        template_content = """
<!DOCTYPE html>
<html>
<head>
    <title>{{filename}}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
        }
        p {
            margin: 0 0 10px;
        }
        b {
            display: block;
            margin: 10px 0;
            font-weight: bold;
        }
    </style>
</head>
<body>
    {% for element in content %}
    {% if element.strip() %}
    {{ element|safe }}
    {% endif %}
    {% endfor %}
</body>
</html>
"""
        template = Template(template_content)
        rendered_content = template.render(content=content, filename=filename)
        
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(rendered_content)

        print(f"Jinja2 template created successfully at {template_path}")
    except Exception as e:
        print(f"An error occurred: {e}")

def remove_toc_anchors(content):
    """Remove instances of {#_TocXXXXX .anchor} from the document."""
    cleaned_content = []
    toc_pattern = re.compile(r'\[\]\{#_Toc\d+ \.anchor\}')

    for line in content:
        cleaned_line = toc_pattern.sub('', line)
        cleaned_content.append(cleaned_line)

    return cleaned_content

def convert_docx_to_markdown(input_file, output_file):
    try:
        # Read the docx file
        doc = Document(input_file)
        markdown_lines = []

        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            # Remove unwanted links and anchors
            line = re.sub(r'\[.*?\]\{#.*?\}', '', line)
            # Remove content inside curly braces {}
            line = re.sub(r'\{.*?\}', '', line)
            if line:
                markdown_lines.append(line)
        
        # Remove TOC anchors from the markdown content
        markdown_lines = remove_toc_anchors(markdown_lines)
        
        # Write the cleaned markdown lines to the output file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(markdown_lines))
        
        print(f"Converted {input_file} to Markdown successfully.")
    except Exception as e:
        print(f"Error converting {input_file} to Markdown: {e}")

def clean_markdown(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # Example of cleaning up the markdown content
        # Remove excessive line breaks
        content = re.sub(r'\n{2,}', '\n\n', content)
        # Remove any other unwanted elements here

        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    except Exception as e:
        print(f"Error cleaning Markdown file {file_path}: {e}")

def convert_markdown_to_html(input_file, output_file):
    try:
        subprocess.run(['pandoc', input_file, '-o', output_file], check=True)
        print(f"Converted {input_file} to HTML successfully.")
    except subprocess.CalledProcessError as e:
        print(f"Error converting {input_file} to HTML: {e}")

def process_files(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith('.pdf'):
            file_path = os.path.join(input_folder, filename)
            content = read_pdf_file(file_path)
            if content:
                output_file_name = f"{os.path.splitext(filename)[0]}.html"
                template_path = os.path.join(output_folder, output_file_name)
                create_jinja2_template(content, template_path, filename)
            else:
                print(f"Failed to read the file: {filename}")
        
        elif filename.endswith('.docx'):
            input_file = os.path.join(input_folder, filename)
            output_file_name = f"{os.path.splitext(filename)[0]}.md"
            output_file = os.path.join(output_folder, output_file_name)
            convert_docx_to_markdown(input_file, output_file)
            clean_markdown(output_file)
            html_output_file = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}.html")
            convert_markdown_to_html(output_file, html_output_file)

input_folder = 'files'
output_folder = 'output'
process_files(input_folder, output_folder)
